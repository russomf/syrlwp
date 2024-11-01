# report_docx.py
# Streamlining Your Research Laboratory with Python
# Authors:   Mark F. Russo, Ph.D and William Neil 
# Publisher: John Wiley & Sons, Inc.
# License:   MIT (https://opensource.org/licenses/MIT)

from io     import BytesIO
from docx   import Document

# Save a matplotlib Figure to a BytesIO and return
def fig2BytesIO(fig):
    bfile = BytesIO()
    fig.savefig(bfile, format='png')
    bfile.seek(0)
    return bfile

# Generate a Word Report
def generate_docx(fig1, fig2, lol, fname):
    # Create a Word Document object
    doc = Document()

    # Add a heading
    doc.add_heading('Screening Report', 0)

    # Add a paragraph with bolded text
    p1 = doc.add_paragraph('This report summarizes results from all dose-response experiments. ')
    r1 = p1.add_run('Consult assay documentation for detailed experimental procedures.')
    r1.bold = True

    # Add Heatmap figure
    doc.add_heading("Experimental Data Heat Map", 2)
    doc.add_picture(fig2BytesIO(fig1))   # Wrap and insert image to document

    # Add IC50 Table
    doc.add_heading("Sample IC50s", 2)

    # Add a table sized to hold all results.
    nrows = len(lol)                    # Get number rows and columns
    ncols = max([len(row) for row in lol])
    tbl = doc.add_table(rows=nrows, cols=ncols)
    tbl.style = 'Light Grid Accent 1'

    # Fill table
    for i in range(nrows):
        tbl.cell(i,0).text = lol[i][0]
        tbl.cell(i,1).text = str(lol[i][1])

    # Add scatter and fits plots
    doc.add_heading("Model Fits", 2)
    doc.add_picture(fig2BytesIO(fig2))   # Wrap and add image to document

    # Save Word document
    doc.save(fname)

if __name__ == '__main__':
    from report_utils import *
    
    # Read data and transpose so columns hold data for single samples
    df = pd.read_csv("data_plate.csv", header=None)
    df = df.transpose()
    
    # Generate heatmap Figure from data
    fig1 = plot_heatmap(df, "Experimental Data")

    # Fit all data to four-parameter logistic model
    guess  = [0, 2, -5, 1]                            # Parameter guesses
    bounds = [[0,0,-9,0], [1,1000,0,1]]               # Parameter bounds
    concs  = [-9, -8, -7, -6, -5, -4, -3, -2, -1, 0]  # log(concentrations)
    param_sets = fit_4PLs(df, concs, guess, bounds)   # Fit all data

    # Create a plot showing all data on eight axes.
    fig2 = plot_fits(df, concs, param_sets=param_sets)
    
    # Collect IC50 results as list of lists
    IC50s = [f'{pow(10, float(pset[2])):0.3E}' for pset in param_sets]
    lol   = [ (f'Sample {i+1}', IC50s[i]) for i in range(len(IC50s)) ]
    lol.insert(0, ["Sample ID", "IC50 [M]"])          # Add column headers
    
    # Generate report Word document
    generate_docx(fig1, fig2, lol, "report.docx")
